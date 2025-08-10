
import os
import json
import gradio as gr
import docx
from docx.shared import RGBColor 
from langchain_google_genai import GoogleGenerativeAI
from langchain_community.vectorstores import FAISS
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain


analysis_prompt_template_str = """
You are the "Corporate Agent," an AI-powered legal assistant for the Abu Dhabi Global Market (ADGM).
Your task is to review a legal document based ONLY on the provided ADGM context.

**ADGM CONTEXT:**
{context}

---

**DOCUMENT TO REVIEW:**
{document_text}

---

**YOUR ANALYSIS:**
Based on the provided context and document, find all legal red flags, errors, or inconsistencies. For each issue, provide:
* "document": The name of the document being analyzed.
* "section": The specific clause or section number (e.g., "Clause 3.1").
* "issue": A clear description of the issue.
* "severity": A severity rating (High, Medium, or Low).
* "suggestion": A suggestion for how to fix it.
* "clause_text": The exact text of the problematic clause.

Present your findings in a structured JSON format with a single key "issues_found", which is a list of objects.
Example JSON:
{{
  "issues_found": [
    {{
      "document": "Articles of Association",
      "section": "Clause 2",
      "issue": "Incorrect jurisdiction. The registered office must be within the ADGM.",
      "severity": "High",
      "suggestion": "Change the jurisdiction to 'Abu Dhabi Global Market'.",
      "clause_text": "The registered office of the Company will be situated in the city of Dubai, United Arab Emirates."
    }}
  ]
}}
"""


COMPANY_INCORPORATION_CHECKLIST = [
    "Application to incorporate a Private Company limited by shares",
    "Articles of Association for a Private Company Limited by Shares",
    "Resolution for Incorporation (LTD - Multiple Shareholders)",
    "Record of Beneficial Owners",
    "ADGM RA Name self declaration"
]

# --- 2. Core Analysis Function ---
def analyze_document_package(api_key, docx_paths):
    
    if not api_key:
        raise gr.Error("Google API Key is required.")
    if not docx_paths:
        raise gr.Error("Please upload at least one .docx file.")

    try:
        
        os.environ["GOOGLE_API_KEY"] = api_key
        llm = GoogleGenerativeAI(model="gemini-1.5-flash", temperature=0.2)
        
       
        identified_types = []
        for doc_path in docx_paths:
            doc = docx.Document(doc_path)
            preview_text = "\n".join([p.text for p in doc.paragraphs[:10]])[:500]
            
            identification_prompt = f"""
            Based on this text preview, what is the most likely document type from the following list?
            List: {COMPANY_INCORPORATION_CHECKLIST}
            
            Preview: "{preview_text}"
            
            Return only the single best matching document type from the list.
            Document Type:
            """
            doc_type = llm.invoke(identification_prompt).strip()
            if doc_type in COMPANY_INCORPORATION_CHECKLIST:
                identified_types.append(doc_type)

       
        missing_docs = [doc for doc in COMPANY_INCORPORATION_CHECKLIST if doc not in identified_types]
        
       
        embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
        vector_store = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
        retriever = vector_store.as_retriever()
        
        main_doc_to_review_path = docx_paths[0]
        doc = docx.Document(main_doc_to_review_path)
        full_text = "\n".join([p.text for p in doc.paragraphs])
        
        relevant_context = retriever.invoke(full_text)
        context_text = "\n".join([d.page_content for d in relevant_context])
        
        analysis_prompt = PromptTemplate(template=analysis_prompt_template_str, input_variables=["context", "document_text"])
        chain = LLMChain(llm=llm, prompt=analysis_prompt)
        llm_response = chain.invoke({"context": context_text, "document_text": full_text})
        
        response_text = llm_response['text'].strip().replace("```json", "").replace("```", "")
        analysis_result = json.loads(response_text)
        
        
        final_report = {
            "process": "Company Incorporation (Assumed)",
            "documents_uploaded": len(identified_types),
            "required_documents": len(COMPANY_INCORPORATION_CHECKLIST),
            "missing_document": missing_docs if missing_docs else "None. All required documents appear to be present.",
            "issues_found": analysis_result.get("issues_found", [])
        }
        
       
        output_docx_path = "reviewed_" + os.path.basename(main_doc_to_review_path)
        all_issues = final_report.get("issues_found", [])
        
        for issue in all_issues:
            clause_text = issue.get("clause_text", "")
            comment_text = f"AI COMMENT: {issue.get('issue')} SUGGESTION: {issue.get('suggestion')}"
            
            for para in doc.paragraphs:
                if clause_text and clause_text in para.text:
                    comment_p = para.insert_paragraph_before(comment_text)
                    comment_p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
                    break 
        
        doc.save(output_docx_path)
        
        return json.dumps(final_report, indent=2), output_docx_path

    except Exception as e:
        raise gr.Error(f"An error occurred: {str(e)}")



with gr.Blocks(theme=gr.themes.Soft()) as app:
    gr.Markdown("# Corporate Agent: ADGM Document Reviewer")
    gr.Markdown("Upload all relevant `.docx` files for your company incorporation package. The agent will check for missing documents and perform a detailed analysis on the first file uploaded.")
    
    with gr.Row():
        # Inputs
        with gr.Column(scale=1):
            api_key_input = gr.Textbox(label="Your Google API Key", type="password")
            # Updated to accept multiple files
            file_input = gr.File(label="Upload ADGM Document Package", file_count="multiple", file_types=[".docx"])
            submit_button = gr.Button("Analyze Document Package", variant="primary")
        
        # Outputs
        with gr.Column(scale=2):
            json_output = gr.JSON(label="Full Analysis Report")
            file_output = gr.File(label="Download Reviewed Document (First File Only)")

    # Connect the button to the new multi-document function
    submit_button.click(
        fn=analyze_document_package,
        inputs=[api_key_input, file_input],
        outputs=[json_output, file_output]
    )


if __name__ == "__main__":
    app.launch(debug=True)
